"""
Project Iceberg web server — bridges the CLI agent to the browser UI.

Serves ui/index.html at / and exposes a JSON REST API at /api/*.

Usage:
    python server.py [--port 5000] [--host 127.0.0.1]
"""

import argparse
import json
import os
import threading
from typing import Optional

import requests as _requests
from flask import Flask, jsonify, request, send_from_directory
from flask.wrappers import Response

from agents.orchestrator_agent import OrchestratorAgent
from tools.registry import TOOL_REGISTRY

# ---------------------------------------------------------------------------
# App setup
# ---------------------------------------------------------------------------

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
UI_DIR = os.path.join(_BASE_DIR, "ui")
CONFIG_PATH = os.path.join(_BASE_DIR, "config.json")
SKILLS_DIR = os.path.join(_BASE_DIR, "Skills")
app = Flask(__name__, static_folder=None)

# ---------------------------------------------------------------------------
# Agent lifecycle — initialised once in a background thread so the server
# can accept requests (and return 503 / retry hints) while MCP servers spin up.
# ---------------------------------------------------------------------------

_agent: Optional[OrchestratorAgent] = None
_agent_lock = threading.Lock()
_agent_ready = threading.Event()
_agent_error: Optional[str] = None


def _boot_agent() -> None:
    global _agent, _agent_error
    try:
        print("[server] booting agent…")
        _agent = OrchestratorAgent(interactive=False)
        print(f"[server] agent ready — LLM: {_agent.llm_info()} | MCP: {_agent.mcp_info()}")
    except Exception as exc:
        _agent_error = str(exc)
        print(f"[server] agent boot failed: {exc}")
    finally:
        _agent_ready.set()


def _require_agent() -> OrchestratorAgent:
    """Block until agent is ready; raise if boot failed."""
    _agent_ready.wait(timeout=90)
    if _agent is None:
        raise RuntimeError(_agent_error or "Agent failed to initialise.")
    return _agent


# ---------------------------------------------------------------------------
# Static files
# ---------------------------------------------------------------------------


@app.route("/")
def index() -> Response:
    return send_from_directory(UI_DIR, "index.html")


# ---------------------------------------------------------------------------
# API
# ---------------------------------------------------------------------------


@app.route("/api/ready")
def api_ready() -> Response:
    """Non-blocking readiness probe for the UI loader."""
    return jsonify({"ready": _agent_ready.is_set(), "error": _agent_error})


@app.route("/api/status")
def api_status() -> Response:
    agent = _require_agent()
    mcp_servers = list(agent.mcp_manager.connected_servers) if agent.mcp_manager else []
    return jsonify(
        {
            "llm": agent.llm_info(),
            "llm_active": agent.llm_active,
            "mcp_servers": mcp_servers,
            "tools_count": len(TOOL_REGISTRY),
            "memory_count": len(agent.long_memory),
            "memory_backend": agent.long_memory.backend_name,
            "history_depth": agent.history_depth(),
        }
    )


@app.route("/api/tools")
def api_tools() -> Response:
    tools = [
        {
            "name": name,
            "description": entry.get("description", ""),
            "category": entry.get("category", "custom"),
            "args": entry.get("args", []),
        }
        for name, entry in sorted(TOOL_REGISTRY.items())
    ]
    return jsonify(tools)


@app.route("/api/servers")
def api_servers() -> Response:
    agent = _require_agent()
    if not agent.mcp_manager:
        return jsonify([])
    result = []
    for server_name in agent.mcp_manager.connected_servers:
        try:
            raw_tools = agent.mcp_manager.list_tools(server_name)
            tool_names = [t.name for t in raw_tools]
        except Exception:
            tool_names = []
        result.append(
            {
                "name": server_name,
                "connected": True,
                "tools": tool_names,
            }
        )
    return jsonify(result)


@app.route("/api/chat", methods=["POST"])
def api_chat() -> Response:
    data = request.get_json(force=True, silent=True) or {}
    message = (data.get("message") or "").strip()
    if not message:
        return jsonify({"error": "empty message"}), 400

    with _agent_lock:
        agent = _require_agent()
        responses = agent.run(message)

    return jsonify({"responses": responses})


@app.route("/api/chat/stream", methods=["POST"])
def api_chat_stream() -> Response:
    """SSE streaming endpoint for the web UI.

    For chat-mode responses the LLM tokens are pushed token-by-token so the UI
    can render text incrementally.  Tool and mixed-mode responses are executed
    synchronously (you can't partially execute a tool) and the final result is
    sent as a single SSE event once ready.

    SSE event format:
        data: {"type": "token",  "text": "..."}   — incremental chat token
        data: {"type": "result", "text": "..."}   — complete tool/mixed result
        data: {"type": "done"}                     — stream finished
        data: {"type": "error",  "text": "..."}   — something went wrong
    """
    import json as _json

    data = request.get_json(force=True, silent=True) or {}
    message = (data.get("message") or "").strip()
    if not message:
        return jsonify({"error": "empty message"}), 400

    def _sse(payload: dict) -> str:
        return f"data: {_json.dumps(payload)}\n\n"

    def _generate():
        try:
            agent = _require_agent()
        except Exception as exc:
            yield _sse({"type": "error", "text": str(exc)})
            yield _sse({"type": "done"})
            return

        # --- classify intent (fast-path avoids LLM for obvious cases) ---
        try:
            classification = agent.dispatcher.classify(message, list(agent._history))
            mode = classification.get("mode", "chat")
        except Exception:
            mode = "chat"

        if mode in ("tools", "mixed"):
            # Tool execution is synchronous — run normally, emit as one result
            try:
                with _agent_lock:
                    responses = agent.run(message)
                full_text = "\n".join(r for r in responses if r)
                yield _sse({"type": "result", "text": full_text})
            except Exception as exc:
                yield _sse({"type": "error", "text": str(exc)})
            yield _sse({"type": "done"})
            return

        # --- pure chat: stream tokens from LLM ---
        if not agent.llm_active or agent.llm is None:
            # No LLM — return the canned no-LLM message
            try:
                with _agent_lock:
                    responses = agent.run(message)
                yield _sse({"type": "result", "text": "\n".join(r for r in responses if r)})
            except Exception as exc:
                yield _sse({"type": "error", "text": str(exc)})
            yield _sse({"type": "done"})
            return

        from agent_core.constants import (LLM_CHAT_MAX_TOKENS,
                                          LLM_CHAT_TEMPERATURE)
        from agents.orchestrator_agent import CHAT_SYSTEM_PROMPT

        try:
            messages_for_llm: list[dict] = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}]
            messages_for_llm.extend(list(agent._history)[-agent._history_limit :])
            messages_for_llm.append({"role": "user", "content": message})

            full_reply = []
            for token in agent.llm.chat_stream(
                messages_for_llm,
                max_tokens=LLM_CHAT_MAX_TOKENS,
                temperature=LLM_CHAT_TEMPERATURE,
            ):
                full_reply.append(token)
                yield _sse({"type": "token", "text": token})

            # Commit the exchange to agent history so follow-ups work.
            # _push() handles the rolling window trim; short_memory expects dicts.
            complete_reply = "".join(full_reply)
            with _agent_lock:
                agent._push("user", message)
                agent._push("assistant", complete_reply)
                agent.short_memory.add(
                    {"input": message, "output": [complete_reply], "mode": "chat"}
                )
                agent.long_memory.add(message)

        except Exception as exc:
            yield _sse({"type": "error", "text": str(exc)})

        yield _sse({"type": "done"})

    return Response(
        _generate(),
        mimetype="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "X-Accel-Buffering": "no",  # disables Nginx buffering if behind a proxy
        },
    )


@app.route("/api/reset", methods=["POST"])
def api_reset() -> Response:
    with _agent_lock:
        agent = _require_agent()
        agent.reset_history()
    return jsonify({"ok": True})


# ---------------------------------------------------------------------------
# LLM management
# ---------------------------------------------------------------------------


def _read_config() -> dict:
    """Read config.json, returning {} on missing / bad JSON."""
    if not os.path.exists(CONFIG_PATH):
        return {}
    try:
        with open(CONFIG_PATH, encoding="utf-8") as f:
            return json.load(f)
    except Exception as exc:
        print(f"[server] config.json parse error: {exc}")
        return {}


def _write_config(updates: dict) -> None:
    """Merge updates into config.json (creates it if missing)."""
    cfg = _read_config()
    cfg.update(updates)
    with open(CONFIG_PATH, "w", encoding="utf-8") as f:
        json.dump(cfg, f, indent=2)


# ---------------------------------------------------------------------------
# Curated Ollama model catalog — shown in the picker even before pulling.
# Organised into two groups: uncensored models and popular general models.
# Sizes are approximate on-disk figures at the default (Q4) quantisation.
# ---------------------------------------------------------------------------

_OLLAMA_CATALOG: list[dict] = [
    # ── Uncensored / minimal-filter ──────────────────────────────────────
    {
        "name": "dolphin3",
        "size": "~5.0 GB",
        "group": "uncensored",
        "desc": "Dolphin 3.0 · Llama 3.1 8B · uncensored",
    },
    {
        "name": "dolphin-mistral",
        "size": "~4.1 GB",
        "group": "uncensored",
        "desc": "Dolphin Mistral 7B · uncensored",
    },
    {
        "name": "dolphin-llama3",
        "size": "~4.9 GB",
        "group": "uncensored",
        "desc": "Dolphin LLaMA 3 8B · uncensored",
    },
    {
        "name": "dolphin-phi",
        "size": "~1.6 GB",
        "group": "uncensored",
        "desc": "Dolphin Phi 2.8B · uncensored · fast",
    },
    {
        "name": "hermes3",
        "size": "~5.0 GB",
        "group": "uncensored",
        "desc": "Hermes 3 · Llama 3.1 8B · uncensored",
    },
    {
        "name": "nous-hermes2",
        "size": "~5.7 GB",
        "group": "uncensored",
        "desc": "Nous Hermes 2 · 10.7B · uncensored",
    },
    {
        "name": "nous-hermes",
        "size": "~7.4 GB",
        "group": "uncensored",
        "desc": "Nous Hermes 13B · uncensored",
    },
    {
        "name": "openhermes",
        "size": "~4.1 GB",
        "group": "uncensored",
        "desc": "OpenHermes 2.5 · Mistral 7B · uncensored",
    },
    {
        "name": "wizard-vicuna-uncensored",
        "size": "~3.8 GB",
        "group": "uncensored",
        "desc": "WizardLM Vicuna 7B · explicitly uncensored",
    },
    {
        "name": "llama2-uncensored",
        "size": "~3.8 GB",
        "group": "uncensored",
        "desc": "LLaMA 2 7B · explicitly uncensored",
    },
    {
        "name": "samantha-mistral",
        "size": "~4.1 GB",
        "group": "uncensored",
        "desc": "Samantha Mistral 7B · uncensored",
    },
    {
        "name": "orca-mini",
        "size": "~2.0 GB",
        "group": "uncensored",
        "desc": "Orca Mini 3B · lightweight · uncensored",
    },
    {
        "name": "neural-chat",
        "size": "~4.1 GB",
        "group": "uncensored",
        "desc": "Neural Chat 7B · Intel · uncensored",
    },
    {
        "name": "stable-beluga",
        "size": "~4.1 GB",
        "group": "uncensored",
        "desc": "Stable Beluga 7B · uncensored",
    },
    {
        "name": "wizardlm-uncensored",
        "size": "~7.4 GB",
        "group": "uncensored",
        "desc": "WizardLM 13B · uncensored",
    },
    # ── Popular general models ───────────────────────────────────────────
    {
        "name": "gemma3:4b",
        "size": "~3.3 GB",
        "group": "popular",
        "desc": "Google Gemma 3 4B · fast · capable",
    },
    {
        "name": "gemma3:12b",
        "size": "~8.1 GB",
        "group": "popular",
        "desc": "Google Gemma 3 12B · best quality on 1080 Ti",
    },
    {
        "name": "qwen2.5:7b",
        "size": "~4.7 GB",
        "group": "popular",
        "desc": "Alibaba Qwen 2.5 7B · strong coding",
    },
    {
        "name": "qwen2.5-coder:7b",
        "size": "~4.7 GB",
        "group": "popular",
        "desc": "Qwen 2.5 Coder 7B · best local coding model",
    },
    {
        "name": "llama3.2:3b",
        "size": "~2.0 GB",
        "group": "popular",
        "desc": "Meta LLaMA 3.2 3B · very fast",
    },
    {
        "name": "llama3.2:8b",
        "size": "~4.9 GB",
        "group": "popular",
        "desc": "Meta LLaMA 3.2 8B · balanced",
    },
    {
        "name": "mistral",
        "size": "~4.1 GB",
        "group": "popular",
        "desc": "Mistral 7B v0.3 · reliable general purpose",
    },
    {
        "name": "deepseek-r1:8b",
        "size": "~4.9 GB",
        "group": "popular",
        "desc": "DeepSeek R1 8B · chain-of-thought reasoning",
    },
    {
        "name": "phi4",
        "size": "~8.9 GB",
        "group": "popular",
        "desc": "Microsoft Phi-4 14B · strong reasoning, fits 1080 Ti",
    },
    {
        "name": "phi3.5",
        "size": "~2.2 GB",
        "group": "popular",
        "desc": "Microsoft Phi-3.5 3.8B · lightweight powerhouse",
    },
    {
        "name": "codellama:7b",
        "size": "~3.8 GB",
        "group": "popular",
        "desc": "Meta Code LLaMA 7B · code generation",
    },
]


@app.route("/api/llm/models")
def api_llm_models() -> Response:
    """Return Ollama models — pulled ones first, then unpulled catalog entries."""
    pulled_models = []
    pulled_names: set[str] = set()

    try:
        r = _requests.get("http://localhost:11434/api/tags", timeout=3)
        if r.ok:
            for m in r.json().get("models", []):
                size_b = m.get("size", 0)
                size_str = f"{size_b / 1e9:.1f} GB" if size_b else "?"
                name = m["name"]
                pulled_names.add(name)
                # Also track the base name (without tag) for catalog dedup
                pulled_names.add(name.split(":")[0])
                pulled_models.append(
                    {
                        "name": name,
                        "size": size_str,
                        "provider": "ollama",
                        "pulled": True,
                        "group": "installed",
                        "desc": "",
                    }
                )
    except Exception:
        pass

    # Append catalog entries that aren't already pulled
    catalog_models = []
    for entry in _OLLAMA_CATALOG:
        base = entry["name"].split(":")[0]
        if entry["name"] not in pulled_names and base not in pulled_names:
            catalog_models.append(
                {
                    "name": entry["name"],
                    "size": entry["size"],
                    "provider": "ollama",
                    "pulled": False,
                    "group": entry["group"],
                    "desc": entry["desc"],
                }
            )

    return jsonify({"models": pulled_models + catalog_models})


@app.route("/api/ollama/pull", methods=["POST"])
def api_ollama_pull() -> Response:
    """Stream an 'ollama pull <model>' operation as SSE.

    Request body: { "model": "dolphin3" }

    SSE events:
        data: {"type": "status",   "text": "pulling manifest"}
        data: {"type": "progress", "text": "downloading…", "pct": 42}
        data: {"type": "done",     "text": "success"}
        data: {"type": "error",    "text": "..."}
    """
    import json as _json

    data = request.get_json(force=True, silent=True) or {}
    model = (data.get("model") or "").strip()
    if not model:
        return jsonify({"error": "model name required"}), 400

    def _sse(payload: dict) -> str:
        return f"data: {_json.dumps(payload)}\n\n"

    def _generate():
        try:
            with _requests.post(
                "http://localhost:11434/api/pull",
                json={"name": model, "stream": True},
                timeout=600,  # large models take a while
                stream=True,
            ) as resp:
                if not resp.ok:
                    yield _sse({"type": "error", "text": f"Ollama returned {resp.status_code}"})
                    return

                for raw_line in resp.iter_lines():
                    if not raw_line:
                        continue
                    try:
                        chunk = _json.loads(raw_line)
                    except Exception:
                        continue

                    status = chunk.get("status", "")
                    total = chunk.get("total", 0)
                    completed = chunk.get("completed", 0)

                    if status == "success":
                        yield _sse({"type": "done", "text": "Pull complete"})
                        return

                    if total and completed:
                        pct = int(completed / total * 100)
                        size_done = f"{completed / 1e9:.1f}"
                        size_total = f"{total / 1e9:.1f}"
                        yield _sse(
                            {
                                "type": "progress",
                                "text": f"{size_done} / {size_total} GB",
                                "pct": pct,
                            }
                        )
                    else:
                        yield _sse({"type": "status", "text": status})

        except Exception as exc:
            yield _sse({"type": "error", "text": str(exc)})

    return Response(
        _generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


@app.route("/api/llm/set", methods=["POST"])
def api_llm_set() -> Response:
    """Hot-swap the active LLM and persist the choice to config.json."""
    from agent_core.dispatcher import Dispatcher
    from agent_core.llm import (AnthropicProvider, LMStudioProvider,
                                OllamaProvider)
    from planner.planner import Planner
    from tools.registry import describe_for_llm

    data = request.get_json(force=True, silent=True) or {}
    provider = (data.get("provider") or "ollama").lower()
    model = (data.get("model") or "").strip()
    host = (data.get("host") or "").strip()
    api_key = (data.get("api_key") or "").strip()

    # Build new provider object
    try:
        if provider == "ollama":
            h = host or "http://localhost:11434"
            new_llm = OllamaProvider(h, model)
        elif provider == "lmstudio":
            h = host or "http://localhost:1234"
            new_llm = LMStudioProvider(h, model)
        elif provider == "anthropic":
            if not api_key:
                api_key = _read_config().get("anthropic_api_key") or ""
            new_llm = AnthropicProvider(api_key, model)
        else:
            return jsonify({"error": f"Unknown provider: {provider}"}), 400
    except Exception as exc:
        return jsonify({"error": str(exc), "needs_restart": True}), 500

    # Hot-swap — replace llm, planner, and dispatcher in-place
    try:
        with _agent_lock:
            agent = _require_agent()
            agent.llm = new_llm
            agent.llm_active = True
            tools_desc = describe_for_llm()
            agent.planner = Planner(llm=new_llm, tools_description=tools_desc)
            agent.dispatcher = Dispatcher(new_llm, tools_description=tools_desc)
        swapped = True
    except Exception as exc:
        print(f"[server] hot-swap failed: {exc}")
        swapped = False

    # Persist to config.json regardless of hot-swap result
    updates: dict = {"llm_provider": provider}
    if provider == "ollama":
        updates["ollama_model"] = model
        if host:
            updates["ollama_host"] = host
    elif provider == "lmstudio":
        updates["lmstudio_model"] = model
        if host:
            updates["lmstudio_host"] = host
    elif provider == "anthropic":
        updates["anthropic_model"] = model
        if api_key:
            updates["anthropic_api_key"] = api_key
    try:
        _write_config(updates)
    except Exception as exc:
        print(f"[server] config write error: {exc}")

    if swapped:
        agent = _require_agent()
        return jsonify({"ok": True, "llm": agent.llm_info(), "needs_restart": False})
    return jsonify(
        {
            "ok": False,
            "needs_restart": True,
            "error": "Hot-swap failed — restart Project Iceberg to apply.",
        }
    )


# ---------------------------------------------------------------------------
# MCP management
# ---------------------------------------------------------------------------


@app.route("/api/mcp/catalog")
def api_mcp_catalog() -> Response:
    """Return the catalog of available MCP servers."""
    from tools.mcp_catalog import get_catalog

    return jsonify({"servers": get_catalog()})


@app.route("/api/mcp/check-npx")
def api_mcp_check_npx() -> Response:
    """
    Check if npx is installed and accessible.

    Returns:
      {
        "installed": true/false,
        "version": "11.12.1" or null
      }
    """
    import subprocess

    try:
        # On Windows, shell=True is needed to resolve .cmd files
        result = subprocess.run(
            ["npx", "--version"],
            capture_output=True,
            text=True,
            timeout=5,
            shell=True,  # Required on Windows for .cmd files
        )

        if result.returncode == 0:
            version = result.stdout.strip()
            return jsonify({"installed": True, "version": version})
        else:
            return jsonify({"installed": False, "version": None})

    except (FileNotFoundError, subprocess.TimeoutExpired):
        return jsonify({"installed": False, "version": None})


@app.route("/api/mcp/add", methods=["POST"])
def api_mcp_add() -> Response:
    """
    Add an MCP server to config.json.

    Request body:
      {
        "server_name": "github",
        "api_key": "optional-key-value"  // For servers that need API keys
      }
    """
    from tools.mcp_catalog import get_server_config

    data = request.get_json(force=True, silent=True) or {}
    server_name = (data.get("server_name") or "").strip()
    api_key = (data.get("api_key") or "").strip()

    if not server_name:
        return jsonify({"error": "server_name required"}), 400

    # Get server config from catalog
    server_entry = get_server_config(server_name)
    if not server_entry:
        return jsonify({"error": f"Unknown server: {server_name}"}), 404

    # Build the MCP server config
    server_config = {
        "command": server_entry["command"],
        "args": server_entry["args"],
        "env": server_entry["env"].copy(),
    }

    # If API key provided and server needs it, replace the placeholder
    if api_key and server_entry.get("requires_api_key"):
        key_field = server_entry.get("api_key_field")
        if key_field:
            server_config["env"][key_field] = api_key

    # Read current config
    cfg = _read_config()

    # Initialize mcp_servers if it doesn't exist
    if "mcp_servers" not in cfg:
        cfg["mcp_servers"] = {}

    # Check if server already exists
    if server_name in cfg["mcp_servers"]:
        return jsonify({"error": f"Server '{server_name}' already exists in config"}), 409

    # Add the server
    cfg["mcp_servers"][server_name] = server_config

    # Write back to config.json
    try:
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(cfg, f, indent=2)
        return jsonify({"ok": True, "server_name": server_name})
    except Exception as exc:
        return jsonify({"error": f"Failed to write config: {exc}"}), 500


# ---------------------------------------------------------------------------
# Skills management
# ---------------------------------------------------------------------------


def _infer_skill_tags(name: str, desc: str) -> list[str]:
    """Infer category tags from skill name and description text."""
    text = (name + " " + desc).lower()
    tags: list[str] = []
    if any(
        w in text
        for w in [
            "ai",
            "llm",
            "agent",
            "model",
            "fine-tun",
            "learning",
            "instinct",
            "framework",
            "cost-aware",
            "pipeline",
        ]
    ):
        tags.append("AI")
    if any(w in text for w in ["data", "sql", "database", "query", "analyt", "statistic"]):
        tags.append("DATA")
    if any(
        w in text
        for w in [
            "devops",
            "deploy",
            "infra",
            "docker",
            "ci/cd",
            "cloud",
            "research",
            "scraping",
            "scrape",
            "headless",
        ]
    ):
        tags.append("OPS")
    if any(
        w in text
        for w in [
            " ui ",
            "ux",
            "frontend",
            "css",
            "canvas",
            "art",
            "visual",
            "figma",
            "react",
            "tailwind",
        ]
    ):
        tags.append("UI")
    if not tags:
        tags.append("ENG")
    return tags


def _parse_skill_meta(skill_dir: str) -> dict:
    """
    Extract display metadata from a skill directory.

    Handles two SKILL.md formats:
      1. YAML front-matter (--- ... ---): reads ``name`` and ``description`` keys.
      2. Plain markdown: uses first H1 as name, first paragraph as description.
    """
    slug = os.path.basename(skill_dir)
    display_name = slug
    desc = ""
    skill_md = os.path.join(skill_dir, "SKILL.md")

    if os.path.exists(skill_md):
        with open(skill_md, encoding="utf-8", errors="ignore") as fh:
            content = fh.read()

        if content.lstrip().startswith("---"):
            # YAML front-matter: extract name and description keys directly
            in_front = False
            for raw_line in content.splitlines():
                line = raw_line.strip()
                if line == "---":
                    if not in_front:
                        in_front = True
                        continue
                    else:
                        break  # end of front-matter
                if not in_front:
                    continue
                if line.startswith("name:") and display_name == slug:
                    display_name = line[len("name:") :].strip().strip("\"'")
                elif line.startswith("description:") and not desc:
                    desc = line[len("description:") :].strip().strip("\"'")[:200]
        else:
            # Plain markdown: first heading = name, first paragraph = description
            for raw_line in content.splitlines():
                line = raw_line.strip()
                if not line:
                    continue
                if line.startswith("#"):
                    if display_name == slug:
                        display_name = line.lstrip("#").strip()
                    continue
                if not desc:
                    desc = line[:200]
                if display_name != slug and desc:
                    break

    return {
        "id": slug,
        "name": display_name,
        "slug": slug,
        "desc": desc or "No description available.",
        "tags": _infer_skill_tags(slug, desc),
        "installed": True,
    }


@app.route("/api/tool/run", methods=["POST"])
def api_tool_run() -> Response:
    """
    Execute a registered tool directly with explicit arguments.

    Bypasses natural-language planning so form-based invocations are exact.
    Safety gates (allow / confirm / deny) still apply.

    Request body:
      { "tool": "scrape_paginated", "args": {"url": "...", "selector": "..."} }

    Response:
      { "ok": true,  "result": "..." }
      { "ok": false, "error": "..." }
    """
    agent = _require_agent()
    if isinstance(agent, Response):
        return agent

    data = request.get_json(force=True) or {}
    tool_name = (data.get("tool") or "").strip()
    args: dict = data.get("args") or {}

    if not tool_name:
        return jsonify({"ok": False, "error": "tool name is required"}), 400

    from agent_core.contracts import ToolCall

    try:
        call = ToolCall(tool=tool_name, args=args)
        results = agent.executor.execute([call])
        if not results:
            return jsonify({"ok": False, "error": "No result returned"})
        r = results[0]
        return jsonify({"ok": r.ok, "result": r.output, "error": r.error})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500


@app.route("/api/skills")
def api_skills() -> Response:
    """
    Return all installed skills found in the Skills/ directory.

    Each skill is a subdirectory containing a SKILL.md file.
    Response: { "skills": [ {id, name, slug, desc, tags, installed}, ... ] }
    """
    if not os.path.isdir(SKILLS_DIR):
        return jsonify({"skills": []})

    skills = []
    try:
        for entry in sorted(os.scandir(SKILLS_DIR), key=lambda e: e.name.lower()):
            if entry.is_dir():
                skills.append(_parse_skill_meta(entry.path))
    except OSError as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({"skills": skills})


@app.route("/api/skills/install", methods=["POST"])
def api_skills_install() -> Response:
    """
    Install a skill by copying a local directory into Skills/.

    Request body: { "path": "/absolute/path/to/skill-dir" }
    Response:     { "ok": true, "name": "skill-name" }
                  { "error": "..." }
    """
    import shutil

    data = request.get_json(force=True) or {}
    src = data.get("path", "").strip()

    if not src:
        return jsonify({"error": "path is required"}), 400
    if not os.path.isdir(src):
        return jsonify({"error": f"Not a directory: {src}"}), 400

    skill_name = os.path.basename(src.rstrip("/\\"))
    if not skill_name:
        return jsonify({"error": "Could not determine skill name from path"}), 400

    dst = os.path.join(SKILLS_DIR, skill_name)
    if os.path.exists(dst):
        return jsonify({"error": f"Skill '{skill_name}' is already installed"}), 409

    try:
        os.makedirs(SKILLS_DIR, exist_ok=True)
        shutil.copytree(src, dst)
        meta = _parse_skill_meta(dst)
        return jsonify({"ok": True, "name": skill_name, "skill": meta})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


# ---------------------------------------------------------------------------
# Vision API
# ---------------------------------------------------------------------------


@app.route("/api/vision/analyze", methods=["POST"])
def api_vision_analyze() -> Response:
    """
    Analyze an uploaded image with the local vision model (Ollama).

    Accepts multipart/form-data:
      - image: the image file (JPEG, PNG, WEBP, etc.)
      - prompt: (optional) question to ask about the image

    Or JSON body:
      { "path": "/absolute/path/to/image.jpg", "prompt": "..." }

    Response:
      { "ok": true,  "result": "Vision model response text" }
      { "ok": false, "error": "..." }
    """
    import tempfile

    from tools.vision_tools import analyze_image

    prompt = "Describe this image in detail."
    tmp_path = None

    # Multipart upload
    if request.files.get("image"):
        f = request.files["image"]
        prompt = request.form.get("prompt", prompt)
        suffix = os.path.splitext(f.filename or ".jpg")[1] or ".jpg"
        try:
            with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
                f.save(tmp)
                tmp_path = tmp.name
            result = analyze_image(tmp_path, prompt)
            return jsonify({"ok": True, "result": result})
        except Exception as exc:
            return jsonify({"ok": False, "error": str(exc)}), 500
        finally:
            if tmp_path and os.path.exists(tmp_path):
                try:
                    os.unlink(tmp_path)
                except Exception:
                    pass

    # JSON body with local path
    data = request.get_json(force=True, silent=True) or {}
    path = data.get("path", "").strip()
    prompt = data.get("prompt", prompt)
    if not path:
        return (
            jsonify(
                {"ok": False, "error": "No image provided. Send 'image' file or 'path' JSON field."}
            ),
            400,
        )

    result = analyze_image(path, prompt)
    return jsonify({"ok": True, "result": result})


@app.route("/api/vision/screen", methods=["POST"])
def api_vision_screen() -> Response:
    """
    Take a screenshot and analyze it with the local vision model.

    Request body (optional):
      { "prompt": "What is happening on screen?" }

    Response:
      { "ok": true, "result": "..." }
    """
    from tools.vision_tools import capture_screen

    data = request.get_json(force=True, silent=True) or {}
    prompt = data.get("prompt", "Describe what is on the screen.")

    result = capture_screen(prompt)
    return jsonify({"ok": "[vision error]" not in result, "result": result})


# ---------------------------------------------------------------------------
# Training / fine-tuning API
# ---------------------------------------------------------------------------


@app.route("/api/training/export", methods=["POST"])
def api_training_export() -> Response:
    """
    Export conversation history to a JSONL fine-tuning dataset.

    Request body (all optional):
      { "format": "alpaca" | "sharegpt", "out": "training/dataset.jsonl" }

    Response:
      { "ok": true, "count": 42, "path": "training/dataset.jsonl" }
    """
    import sys

    data = request.get_json(force=True, silent=True) or {}
    fmt = data.get("format", "alpaca")
    out = data.get("out", os.path.join(_BASE_DIR, "training", "dataset.jsonl"))

    # Import relative to the project root
    sys.path.insert(0, _BASE_DIR)
    try:
        from training.export_history import export  # type: ignore

        count = export(
            log_dir=os.path.join(_BASE_DIR, "logs"),
            memory_path=os.path.join(_BASE_DIR, "memory_store.json"),
            output_path=out,
            fmt=fmt,
        )
        return jsonify({"ok": count > 0, "count": count, "path": out})
    except Exception as exc:
        return jsonify({"ok": False, "error": str(exc)}), 500


# ---------------------------------------------------------------------------
# File Viewer API
# ---------------------------------------------------------------------------

_UPLOAD_DIR = os.path.join(_BASE_DIR, "uploads")
os.makedirs(_UPLOAD_DIR, exist_ok=True)


def _read_file_for_api(file_path: str, sheet: str = "", max_rows: int = 1000) -> dict:
    """
    Read a file and return structured JSON data for the UI viewer.

    Returns:
      {
        "type":    "table" | "text" | "error",
        "name":    filename,
        "ext":     extension,
        "size":    bytes,
        "sheets":  [...] (Excel only),
        "sheet":   active sheet name (Excel only),
        "headers": [...] (table only),
        "rows":    [[...], ...] (table only),
        "content": "..." (text only),
        "truncated": bool,
      }
    """
    import csv
    import json as _json
    from pathlib import Path

    path = Path(file_path)
    if not path.exists():
        return {"type": "error", "message": f"File not found: {file_path}"}
    if not path.is_file():
        return {"type": "error", "message": f"Not a file: {file_path}"}

    ext = path.suffix.lower()
    size = path.stat().st_size
    base = {"name": path.name, "ext": ext, "size": size}

    # ── EXCEL ──────────────────────────────────────────────
    if ext in (".xlsx", ".xls", ".xlsm", ".ods"):
        try:
            import openpyxl

            wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
            sheets = wb.sheetnames
            ws = wb[sheet] if sheet and sheet in sheets else wb.active
            active = ws.title

            rows = []
            for row in ws.iter_rows(values_only=True):
                rows.append([str(cell) if cell is not None else "" for cell in row])
            wb.close()

            # Trim trailing empty rows
            while rows and all(c == "" for c in rows[-1]):
                rows.pop()

            headers = rows[0] if rows else []
            data_rows = rows[1 : max_rows + 1]
            truncated = len(rows) - 1 > max_rows

            return {
                **base,
                "type": "table",
                "sheets": sheets,
                "sheet": active,
                "headers": headers,
                "rows": data_rows,
                "total_rows": len(rows) - 1,
                "truncated": truncated,
            }
        except Exception as e:
            return {"type": "error", "message": str(e)}

    # ── CSV ────────────────────────────────────────────────
    elif ext == ".csv":
        try:
            with open(path, newline="", encoding="utf-8-sig", errors="replace") as f:
                sample = f.read(4096)
                f.seek(0)
                try:
                    dialect = csv.Sniffer().sniff(sample)
                except csv.Error:
                    dialect = csv.excel
                reader = csv.reader(f, dialect)
                rows = list(reader)

            while rows and all(c.strip() == "" for c in rows[-1]):
                rows.pop()

            headers = rows[0] if rows else []
            data_rows = rows[1 : max_rows + 1]
            truncated = len(rows) - 1 > max_rows

            return {
                **base,
                "type": "table",
                "sheets": [],
                "sheet": "",
                "headers": headers,
                "rows": data_rows,
                "total_rows": len(rows) - 1,
                "truncated": truncated,
            }
        except Exception as e:
            return {"type": "error", "message": str(e)}

    # ── JSON ───────────────────────────────────────────────
    elif ext in (".json", ".jsonl"):
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
            truncated = False
            if ext == ".jsonl":
                lines = [l for l in text.splitlines() if l.strip()]
                content = "\n".join(lines[:200])
                truncated = len(lines) > 200
            else:
                try:
                    parsed = _json.loads(text)
                    content = _json.dumps(parsed, indent=2, ensure_ascii=False)
                except Exception:
                    content = text
                if len(content) > 100000:
                    content = content[:100000]
                    truncated = True
            return {
                **base,
                "type": "text",
                "language": "json",
                "content": content,
                "truncated": truncated,
            }
        except Exception as e:
            return {"type": "error", "message": str(e)}

    # ── PDF ────────────────────────────────────────────────
    elif ext == ".pdf":
        try:
            import pdfplumber

            pages_text = []
            with pdfplumber.open(path) as pdf:
                total_pages = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    pages_text.append(f"--- Page {i+1} ---\n{page.extract_text() or ''}")
            content = "\n\n".join(pages_text)
            truncated = False
            if len(content) > 100000:
                content = content[:100000]
                truncated = True
            return {
                **base,
                "type": "text",
                "language": "text",
                "content": content,
                "pages": total_pages,
                "truncated": truncated,
            }
        except ImportError:
            return {"type": "error", "message": "pdfplumber not installed: pip install pdfplumber"}
        except Exception as e:
            return {"type": "error", "message": str(e)}

    # ── DOCX ───────────────────────────────────────────────
    elif ext == ".docx":
        try:
            from docx import Document

            doc = Document(path)
            parts = []
            for para in doc.paragraphs:
                t = para.text.strip()
                if not t:
                    continue
                style = para.style.name if para.style else ""
                if "Heading" in style:
                    lvl = style.replace("Heading", "").strip()
                    prefix = "#" * (int(lvl) if lvl.isdigit() else 1)
                    parts.append(f"{prefix} {t}")
                else:
                    parts.append(t)
            for table in doc.tables:
                for row in table.rows:
                    parts.append(" | ".join(c.text.strip() for c in row.cells))
            content = "\n".join(parts)
            truncated = False
            if len(content) > 100000:
                content = content[:100000]
                truncated = True
            return {
                **base,
                "type": "text",
                "language": "markdown",
                "content": content,
                "truncated": truncated,
            }
        except ImportError:
            return {
                "type": "error",
                "message": "python-docx not installed: pip install python-docx",
            }
        except Exception as e:
            return {"type": "error", "message": str(e)}

    # ── PLAIN TEXT / CODE ──────────────────────────────────
    else:
        try:
            text = path.read_text(encoding="utf-8", errors="replace")
        except Exception as e:
            return {"type": "error", "message": str(e)}

        lang_map = {
            ".py": "python",
            ".js": "javascript",
            ".ts": "typescript",
            ".html": "html",
            ".css": "css",
            ".sql": "sql",
            ".json": "json",
            ".yaml": "yaml",
            ".yml": "yaml",
            ".md": "markdown",
            ".sh": "bash",
            ".ps1": "powershell",
        }
        lang = lang_map.get(ext, "text")
        truncated = len(text) > 100000
        return {
            **base,
            "type": "text",
            "language": lang,
            "content": text[:100000],
            "truncated": truncated,
        }


@app.route("/api/file/view", methods=["POST"])
def api_file_view() -> Response:
    """
    Read a file by path and return structured data for the viewer.

    Body: { "path": "/full/path/to/file.xlsx", "sheet": "Sheet1", "max_rows": 1000 }
    """
    body = request.get_json(silent=True) or {}
    file_path = body.get("path", "").strip()
    sheet = body.get("sheet", "")
    max_rows = int(body.get("max_rows", 1000))

    if not file_path:
        return jsonify({"type": "error", "message": "No file path provided."}), 400

    result = _read_file_for_api(file_path, sheet=sheet, max_rows=max_rows)
    return jsonify(result)


@app.route("/api/file/upload", methods=["POST"])
def api_file_upload() -> Response:
    """
    Accept a multipart file upload, save it to the uploads directory,
    and return the structured file data for the viewer.

    Form fields: file (required), sheet (optional), max_rows (optional)
    """
    if "file" not in request.files:
        return jsonify({"type": "error", "message": "No file in request."}), 400

    f = request.files["file"]
    if not f.filename:
        return jsonify({"type": "error", "message": "Empty filename."}), 400

    import werkzeug.utils

    safe_name = werkzeug.utils.secure_filename(f.filename)
    save_path = os.path.join(_UPLOAD_DIR, safe_name)
    f.save(save_path)

    sheet = request.form.get("sheet", "")
    max_rows = int(request.form.get("max_rows", 1000))

    result = _read_file_for_api(save_path, sheet=sheet, max_rows=max_rows)
    result["upload_path"] = save_path
    return jsonify(result)


# ── File tree ──────────────────────────────────────────────────────────────


@app.route("/api/files/tree")
def api_files_tree() -> Response:
    """Return a recursive file/folder tree of the project directory."""
    import fnmatch

    IGNORE = {
        "__pycache__",
        ".git",
        ".venv",
        "venv",
        "node_modules",
        "*.pyc",
        "*.pyo",
        "*.egg-info",
        ".pytest_cache",
        "*.log",
        "uploads",
        "logs",
        ".idea",
        ".vscode",
    }

    def _should_skip(name: str) -> bool:
        for pat in IGNORE:
            if fnmatch.fnmatch(name, pat):
                return True
        return False

    def _walk(path: str, depth: int = 0) -> list:
        if depth > 6:
            return []
        entries = []
        try:
            items = sorted(
                os.listdir(path),
                key=lambda x: (not os.path.isdir(os.path.join(path, x)), x.lower()),
            )
            for name in items:
                if _should_skip(name):
                    continue
                full = os.path.join(path, name)
                is_dir = os.path.isdir(full)
                entry = {
                    "name": name,
                    "path": full,
                    "type": "dir" if is_dir else "file",
                    "ext": os.path.splitext(name)[1].lower() if not is_dir else "",
                }
                if is_dir:
                    entry["children"] = _walk(full, depth + 1)
                entries.append(entry)
        except PermissionError:
            pass
        return entries

    tree = _walk(_BASE_DIR)
    return jsonify({"root": _BASE_DIR, "tree": tree})


# ── Terminal ───────────────────────────────────────────────────────────────


@app.route("/api/terminal/run", methods=["POST"])
def api_terminal_run() -> Response:
    """
    Run a shell command and stream output via SSE.

    Request body: { "command": "dir", "shell": "cmd" | "powershell", "cwd": "..." }
    SSE events:
        data: {"type": "stdout", "text": "..."}
        data: {"type": "stderr", "text": "..."}
        data: {"type": "done",   "code": 0}
        data: {"type": "error",  "text": "..."}
    """
    import json as _json
    import subprocess

    data = request.get_json(force=True, silent=True) or {}
    command = (data.get("command") or "").strip()
    shell = (data.get("shell") or "cmd").lower()
    cwd = (data.get("cwd") or _BASE_DIR).strip()

    if not command:
        return jsonify({"error": "command required"}), 400

    # Safety: block obviously destructive commands
    _BLOCKED = ["format ", "del /f /s /q c:\\", "rm -rf /", ":(){ :|:& };:"]
    for blocked in _BLOCKED:
        if blocked in command.lower():
            return jsonify({"error": f"Command blocked for safety: {command}"}), 403

    def _sse(payload: dict) -> str:
        return f"data: {_json.dumps(payload)}\n\n"

    def _generate():
        if shell == "powershell":
            cmd_args = ["powershell", "-NoProfile", "-NonInteractive", "-Command", command]
        else:
            cmd_args = ["cmd", "/c", command]

        try:
            proc = subprocess.Popen(
                cmd_args,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                cwd=cwd if os.path.isdir(cwd) else _BASE_DIR,
                text=True,
                encoding="utf-8",
                errors="replace",
            )

            # Stream stdout line-by-line
            def _read_stderr():
                for line in proc.stderr:
                    pass  # collected below via communicate

            # Use communicate with timeout to avoid hanging
            try:
                stdout, stderr = proc.communicate(timeout=30)
            except subprocess.TimeoutExpired:
                proc.kill()
                stdout, stderr = proc.communicate()
                yield _sse({"type": "stderr", "text": "[Command timed out after 30s]"})

            if stdout:
                for line in stdout.splitlines(keepends=True):
                    yield _sse({"type": "stdout", "text": line.rstrip("\n")})
            if stderr:
                for line in stderr.splitlines(keepends=True):
                    yield _sse({"type": "stderr", "text": line.rstrip("\n")})

            yield _sse({"type": "done", "code": proc.returncode})

        except FileNotFoundError as e:
            yield _sse({"type": "error", "text": f"Shell not found: {e}"})
        except Exception as e:
            yield _sse({"type": "error", "text": str(e)})

    return Response(
        _generate(),
        mimetype="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── Prompts ────────────────────────────────────────────────────────────────

_DATA_DIR = os.path.join(_BASE_DIR, "data")
_PROMPTS_FILE = os.path.join(_DATA_DIR, "prompts.json")
os.makedirs(_DATA_DIR, exist_ok=True)


def _load_prompts() -> list:
    if not os.path.isfile(_PROMPTS_FILE):
        return []
    try:
        with open(_PROMPTS_FILE, encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []


def _save_prompts(prompts: list) -> None:
    with open(_PROMPTS_FILE, "w", encoding="utf-8") as f:
        json.dump(prompts, f, indent=2, ensure_ascii=False)


@app.route("/api/prompts", methods=["GET"])
def api_prompts_list() -> Response:
    return jsonify({"prompts": _load_prompts()})


@app.route("/api/prompts", methods=["POST"])
def api_prompts_create() -> Response:
    import uuid

    data = request.get_json(force=True, silent=True) or {}
    name = (data.get("name") or "Untitled prompt").strip()
    text = (data.get("text") or "").strip()
    if not text:
        return jsonify({"error": "text required"}), 400
    prompts = _load_prompts()
    new_prompt = {
        "id": str(uuid.uuid4()),
        "name": name,
        "text": text,
        "enabled": data.get("enabled", True),
    }
    prompts.append(new_prompt)
    _save_prompts(prompts)
    return jsonify({"prompt": new_prompt}), 201


@app.route("/api/prompts/<prompt_id>", methods=["PUT"])
def api_prompts_update(prompt_id: str) -> Response:
    data = request.get_json(force=True, silent=True) or {}
    prompts = _load_prompts()
    for p in prompts:
        if p["id"] == prompt_id:
            if "name" in data:
                p["name"] = data["name"]
            if "text" in data:
                p["text"] = data["text"]
            if "enabled" in data:
                p["enabled"] = bool(data["enabled"])
            _save_prompts(prompts)
            return jsonify({"prompt": p})
    return jsonify({"error": "not found"}), 404


@app.route("/api/prompts/<prompt_id>", methods=["DELETE"])
def api_prompts_delete(prompt_id: str) -> Response:
    prompts = _load_prompts()
    prompts = [p for p in prompts if p["id"] != prompt_id]
    _save_prompts(prompts)
    return jsonify({"ok": True})


# ── Conversation history ───────────────────────────────────────────────────


@app.route("/api/history")
def api_history() -> Response:
    """Return the current conversation history from the agent's short memory."""
    try:
        agent = _require_agent()
        history = []
        if hasattr(agent, "memory") and hasattr(agent.memory, "get_history"):
            history = agent.memory.get_history()
        elif hasattr(agent, "short_memory"):
            history = list(agent.short_memory.history)
        return jsonify({"history": history})
    except Exception as e:
        return jsonify({"history": [], "error": str(e)})


@app.route("/api/restart", methods=["POST"])
def api_restart() -> Response:
    """
    Restart Project Iceberg server.

    This endpoint triggers a clean shutdown and restart of the Flask app.
    """
    import sys

    def restart_server() -> None:
        """Restart the server by re-executing the current script."""
        import time

        time.sleep(0.5)  # Give time for response to send
        python = sys.executable
        os.execl(python, python, *sys.argv)

    threading.Thread(target=restart_server, daemon=True).start()
    return jsonify({"ok": True, "message": "Restarting Project Iceberg..."})


# ---------------------------------------------------------------------------
# Self-improvement: lessons API
# ---------------------------------------------------------------------------


@app.route("/api/lessons", methods=["GET"])
def api_lessons_list() -> Response:
    """Return all stored self-improvement lessons, newest first."""
    try:
        agent = _require_agent()
        store = getattr(agent, "lesson_store", None)
        if store is None:
            return jsonify({"lessons": [], "error": "lesson store not initialised"})
        lessons = [l.to_dict() for l in store.all()]
        return jsonify({"lessons": lessons, "count": len(lessons)})
    except Exception as e:
        return jsonify({"lessons": [], "error": str(e)})


@app.route("/api/lessons/<lesson_id>", methods=["DELETE"])
def api_lessons_delete(lesson_id: str) -> Response:
    """Delete a specific lesson by ID."""
    try:
        agent = _require_agent()
        store = getattr(agent, "lesson_store", None)
        if store is None:
            return jsonify({"ok": False, "error": "lesson store not initialised"})
        removed = store.remove(lesson_id)
        if removed:
            return jsonify({"ok": True})
        return jsonify({"ok": False, "error": "lesson not found"}), 404
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


@app.route("/api/lessons/clear", methods=["POST"])
def api_lessons_clear() -> Response:
    """Delete all stored lessons (fresh start)."""
    try:
        agent = _require_agent()
        store = getattr(agent, "lesson_store", None)
        if store is None:
            return jsonify({"ok": False, "error": "lesson store not initialised"})
        all_lessons = store.all()
        for lesson in all_lessons:
            store.remove(lesson.id)
        return jsonify({"ok": True, "deleted": len(all_lessons)})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# Skill engine — match query against indexed skills
# ---------------------------------------------------------------------------


@app.route("/api/skills/match", methods=["GET"])
def api_skills_match() -> Response:
    """
    Return skills most relevant to a query string.

    GET /api/skills/match?q=write+a+python+script&top=5

    Response: {"matches": [{"slug":…,"name":…,"description":…,"score":…}], "count":…}
    """
    query   = request.args.get("q", "").strip()
    top_n   = min(int(request.args.get("top", 5)), 10)

    if not query:
        return jsonify({"matches": [], "count": 0})

    try:
        agent  = _require_agent()
        engine = getattr(agent, "skill_engine", None)
        if engine is None:
            return jsonify({"matches": [], "count": 0, "error": "skill engine not initialised"})

        matches = engine.match(query, top_n=top_n)
        return jsonify({
            "matches": [
                {
                    "slug"        : m.slug,
                    "name"        : m.name,
                    "description" : m.description,
                    "score"       : m.score,
                }
                for m in matches
            ],
            "count"      : len(matches),
            "total_indexed": engine.skill_count(),
        })
    except Exception as e:
        return jsonify({"matches": [], "count": 0, "error": str(e)}), 500


@app.route("/api/skills/rebuild", methods=["POST"])
def api_skills_rebuild() -> Response:
    """Force a full rebuild of the skill index from disk."""
    try:
        agent  = _require_agent()
        engine = getattr(agent, "skill_engine", None)
        if engine is None:
            return jsonify({"ok": False, "error": "skill engine not initialised"})
        count = engine.rebuild()
        return jsonify({"ok": True, "count": count})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# Native file / folder dialog  (tkinter subprocess, Windows-safe)
# ---------------------------------------------------------------------------


def _run_tkinter_dialog(script: str) -> str:
    """Run a tiny tkinter script in a subprocess and return the chosen path."""
    import subprocess as _sp
    import sys as _sys

    result = _sp.run(
        [_sys.executable, "-c", script],
        capture_output=True,
        text=True,
        timeout=120,  # 2-minute timeout — user has time to browse
    )
    return result.stdout.strip()


@app.route("/api/dialog/open-file")
def api_dialog_open_file() -> Response:
    """Show a native open-file dialog; return the selected path."""
    title = request.args.get("title", "Open File")
    script = (
        "import tkinter as tk\n"
        "from tkinter import filedialog\n"
        "root = tk.Tk(); root.withdraw()\n"
        "root.wm_attributes('-topmost', 1)\n"
        f"path = filedialog.askopenfilename(title={title!r})\n"
        "root.destroy()\n"
        "print(path)\n"
    )
    try:
        path = _run_tkinter_dialog(script)
        return jsonify({"path": path, "cancelled": not bool(path)})
    except Exception as e:
        return jsonify({"path": "", "cancelled": True, "error": str(e)}), 500


@app.route("/api/dialog/save-file")
def api_dialog_save_file() -> Response:
    """Show a native save-as dialog; return the selected path."""
    title = request.args.get("title", "Save File")
    script = (
        "import tkinter as tk\n"
        "from tkinter import filedialog\n"
        "root = tk.Tk(); root.withdraw()\n"
        "root.wm_attributes('-topmost', 1)\n"
        f"path = filedialog.asksaveasfilename(title={title!r})\n"
        "root.destroy()\n"
        "print(path)\n"
    )
    try:
        path = _run_tkinter_dialog(script)
        return jsonify({"path": path, "cancelled": not bool(path)})
    except Exception as e:
        return jsonify({"path": "", "cancelled": True, "error": str(e)}), 500


@app.route("/api/dialog/open-folder")
def api_dialog_open_folder() -> Response:
    """Show a native folder picker; return the selected path."""
    title = request.args.get("title", "Select Folder")
    script = (
        "import tkinter as tk\n"
        "from tkinter import filedialog\n"
        "root = tk.Tk(); root.withdraw()\n"
        "root.wm_attributes('-topmost', 1)\n"
        f"path = filedialog.askdirectory(title={title!r})\n"
        "root.destroy()\n"
        "print(path)\n"
    )
    try:
        path = _run_tkinter_dialog(script)
        return jsonify({"path": path, "cancelled": not bool(path)})
    except Exception as e:
        return jsonify({"path": "", "cancelled": True, "error": str(e)}), 500


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------


def main() -> None:
    parser = argparse.ArgumentParser(description="Project Iceberg web server")
    parser.add_argument("--port", type=int, default=5000)
    parser.add_argument("--host", default="127.0.0.1")
    args = parser.parse_args()

    threading.Thread(target=_boot_agent, daemon=True).start()
    app.run(host=args.host, port=args.port, debug=False)


if __name__ == "__main__":
    main()
